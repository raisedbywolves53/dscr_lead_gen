"""
DSCR Lead Gen — Compliance One-Pager (Word Document)
=====================================================

Professional compliance document for LO prospects and their compliance officers.
Addresses "is this legal?" objection with citations and industry precedent.

Usage:
    python scripts/build_compliance_sheet.py
    python scripts/build_compliance_sheet.py --output data/workbooks/compliance.docx
"""

import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent
LOGO_PATH = PROJECT_DIR / "assets" / "logo.png"

# ── Colors ────────────────────────────────────────────────────────────
NAVY = RGBColor(0x0A, 0x23, 0x42)
TEAL = RGBColor(0x1A, 0x6B, 0x6A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x21, 0x21, 0x21)
DARK_GRAY = RGBColor(0x59, 0x59, 0x59)
MED_GRAY = RGBColor(0x99, 0x99, 0x99)
GREEN = RGBColor(0x22, 0x8B, 0x22)
AMBER = RGBColor(0xB4, 0x78, 0x00)
ACCENT = RGBColor(0x00, 0x66, 0xB3)


def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def add_formatted_run(paragraph, text, bold=False, italic=False, size=10,
                      color=BLACK, font_name="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font_name
    return run


def add_section_heading(doc, text):
    """Add a navy section heading with teal bottom border."""
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"
    # Bottom border via paragraph formatting
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="2" w:color="1A6B6A"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_bullet(doc, text, bold_prefix=None, size=9):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    p.space_after = Pt(2)
    p.space_before = Pt(1)
    if bold_prefix:
        add_formatted_run(p, bold_prefix, bold=True, size=size)
        add_formatted_run(p, text, size=size)
    else:
        add_formatted_run(p, text, size=size)
    return p


def add_body(doc, text, size=9, space_after=4):
    p = doc.add_paragraph()
    p.space_after = Pt(space_after)
    p.space_before = Pt(1)
    add_formatted_run(p, text, size=size)
    return p


def build_compliance_doc(output_path):
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    # ── Default font ──────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)
    style.font.color.rgb = BLACK

    # ════════════════════════════════════════════════════════════
    # HEADER BANNER (table hack for colored background)
    # ════════════════════════════════════════════════════════════
    # Light header background so the teal logo pops
    HEADER_BG = "EDF3F3"  # very light teal-gray — complements the logo

    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in header_table.rows:
        for cell in row.cells:
            set_cell_shading(cell, HEADER_BG)
            set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
            # Clean borders: none on sides, teal bottom
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="1A6B6A"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)

    # Left cell: title + subtitle in navy/teal on light bg
    left_cell = header_table.cell(0, 0)
    left_cell.width = Inches(5.2)
    p = left_cell.paragraphs[0]
    p.space_after = Pt(0)
    add_formatted_run(p, "Data Compliance & Legal Framework", bold=True, size=16, color=NAVY)
    p2 = left_cell.add_paragraph()
    p2.space_before = Pt(2)
    p2.space_after = Pt(0)
    add_formatted_run(p2, "DSCR Investor Intelligence  |  Public Records + Contact Enrichment",
                      size=8, color=TEAL)

    # Right cell: logo
    right_cell = header_table.cell(0, 1)
    right_cell.width = Inches(1.8)
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.space_after = Pt(0)
    if LOGO_PATH.exists():
        try:
            run = rp.add_run()
            run.add_picture(str(LOGO_PATH), height=Inches(0.75))
        except Exception:
            add_formatted_run(rp, "Still Mind Creative", bold=True, size=11, color=TEAL)
    else:
        add_formatted_run(rp, "Still Mind Creative", bold=True, size=11, color=TEAL)

    # ════════════════════════════════════════════════════════════
    # EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════
    doc.add_paragraph()  # spacer

    summary_table = doc.add_table(rows=1, cols=1)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sc = summary_table.cell(0, 0)
    set_cell_shading(sc, "F0F8F7")
    set_cell_margins(sc, top=100, bottom=100, left=160, right=160)
    # Left border accent
    tcPr = sc._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="18" w:space="0" w:color="1A6B6A"/>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D0D8D8"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D8D8"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="D0D8D8"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    sp = sc.paragraphs[0]
    sp.space_after = Pt(4)
    add_formatted_run(sp, "EXECUTIVE SUMMARY", bold=True, size=9, color=NAVY)

    sp2 = sc.add_paragraph()
    sp2.space_after = Pt(0)
    add_formatted_run(sp2,
        "This product is built entirely from publicly available government records and commercially licensed "
        "contact data. It does not use credit bureau data, trigger leads, or any FCRA-regulated consumer "
        "reports. This is the same category of data product offered by the largest property data companies "
        "in the United States -- an established, legally validated model in a multi-billion-dollar industry.",
        size=9)

    # ════════════════════════════════════════════════════════════
    # SECTION 1: DATA SOURCES
    # ════════════════════════════════════════════════════════════
    add_section_heading(doc, "Data Categories & Legal Basis")

    sources = [
        ("Government Property Records", "COMPLIANT",
         "Public by law in all 50 states under Freedom of Information and Public Records Acts. "
         "Facts contained in government records are not copyrightable (Feist v. Rural Telephone, 499 U.S. 340)."),
        ("State Business Filings", "COMPLIANT",
         "Corporate and LLC registrations are public notice records maintained by every state. "
         "Routinely used by title companies, law firms, and due diligence providers."),
        ("Federal Public Databases", "COMPLIANT",
         "Government-published data available through official public APIs and bulk downloads. "
         "No restrictions on commercial use of federal public data."),
        ("Contact Enrichment", "COMPLIANT",
         "Commercially licensed contact data sourced from public records and data cooperatives. "
         "Sold for marketing outreach -- not a consumer report under FCRA."),
    ]

    src_table = doc.add_table(rows=len(sources) + 1, cols=3)
    src_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for ci, hdr in enumerate(["Data Category", "Status", "Legal Basis"]):
        cell = src_table.cell(0, ci)
        set_cell_shading(cell, "0A2342")
        set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
        p = cell.paragraphs[0]
        p.space_after = Pt(0)
        add_formatted_run(p, hdr, bold=True, size=8, color=WHITE)

    widths = [Inches(1.5), Inches(0.9), Inches(4.5)]
    for ri, (source, status, basis) in enumerate(sources, 1):
        bg = "F8F9FB" if ri % 2 == 0 else "FFFFFF"
        for ci, (text, bold, color) in enumerate([
            (source, True, BLACK),
            (status, True, GREEN),
            (basis, False, DARK_GRAY),
        ]):
            cell = src_table.cell(ri, ci)
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=30, bottom=30, left=80, right=80)
            cell.width = widths[ci]
            p = cell.paragraphs[0]
            p.space_after = Pt(0)
            add_formatted_run(p, text, bold=bold, size=8, color=color)

    # ════════════════════════════════════════════════════════════
    # SECTION 2: WHAT THIS IS NOT
    # ════════════════════════════════════════════════════════════
    add_section_heading(doc, "What This Product Is NOT")

    nots = [
        ("NOT trigger leads", " -- trigger leads use credit bureau inquiry data sold to competing lenders, now banned by HPPA (2026). This product does not touch credit data."),
        ("NOT a consumer report", " -- contains no credit scores, payment history, account balances, or credit inquiries."),
        ("NOT FCRA-regulated", " -- data is compiled for marketing outreach, not credit, employment, or insurance decisions."),
        ("NOT a referral fee", " -- flat per-lead pricing with no per-closed-loan fees. RESPA Section 8(c)(2) safe harbor applies."),
    ]
    for bold_part, rest in nots:
        add_bullet(doc, rest, bold_prefix=bold_part, size=9)

    # ════════════════════════════════════════════════════════════
    # SECTION 3: FCRA POSITION
    # ════════════════════════════════════════════════════════════
    add_section_heading(doc, "FCRA Position (15 USC 1681)")

    add_body(doc,
        'Under 15 USC 1681a(d), a "consumer report" must bear on creditworthiness and be used for '
        "credit, employment, or insurance decisions. This product contains public property data and contact "
        "information sold for marketing purposes only. It is not assembled or evaluated for credit decisions. "
        "Every delivery includes an FCRA disclaimer and prohibited-use certification, consistent with "
        "industry standard practice.", size=9)

    # Industry precedent box
    prec_table = doc.add_table(rows=1, cols=1)
    prec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    pc = prec_table.cell(0, 0)
    set_cell_shading(pc, "EDF1F7")
    set_cell_margins(pc, top=80, bottom=80, left=140, right=140)
    pp = pc.paragraphs[0]
    pp.space_after = Pt(3)
    add_formatted_run(pp, "INDUSTRY PRECEDENT:  ", bold=True, size=8, color=NAVY)
    add_formatted_run(pp,
        "CoreLogic ($6B+ revenue), ATTOM Data, PropStream, BatchData, and Reonomy all operate "
        "this identical model: aggregate public property records, enrich with contact data, sell to "
        "real estate and mortgage professionals. All maintain explicit FCRA disclaimers and prohibited-use "
        "policies in their terms of service.",
        size=8, color=DARK_GRAY)

    # ════════════════════════════════════════════════════════════
    # SECTION 4: OUTREACH COMPLIANCE
    # ════════════════════════════════════════════════════════════
    add_section_heading(doc, "Outreach Compliance: Who Owns What")

    outreach = [
        ("DNC Scrubbing", "WE PROVIDE",
         "Leads are delivered with DNC flags and phone type classification. Loan officer should verify before dialing."),
        ("TCPA (Cold Calls)", "LOAN OFFICER",
         "Manual calls to cell phones are legal with DNC compliance. No autodialers to cell phones without written consent."),
        ("Phone Type ID", "WE PROVIDE",
         "Every number classified as Mobile / Landline / VoIP so the loan officer knows which TCPA rules apply per contact."),
        ("CAN-SPAM (Email)", "LOAN OFFICER",
         "Cold email is legal. Must include: physical address, unsubscribe link, honest subject line, honor opt-outs in 10 days."),
        ("Calling Hours", "LOAN OFFICER",
         "Federal: 8am-9pm recipient's time zone. Florida: 8am-8pm. Max 3 calls per 24 hours per person in FL."),
        ("State Licensing", "LOAN OFFICER",
         "Must hold active NMLS license in borrower's state. FL independent LOs may need a telemarketing license for cold calls."),
    ]

    out_table = doc.add_table(rows=len(outreach) + 1, cols=3)
    out_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ci, hdr in enumerate(["Requirement", "Responsibility", "Details"]):
        cell = out_table.cell(0, ci)
        set_cell_shading(cell, "0A2342")
        set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
        p = cell.paragraphs[0]
        p.space_after = Pt(0)
        add_formatted_run(p, hdr, bold=True, size=8, color=WHITE)

    for ri, (req, resp, detail) in enumerate(outreach, 1):
        bg = "F8F9FB" if ri % 2 == 0 else "FFFFFF"
        resp_color = ACCENT if "LOAN" in resp else GREEN

        for ci, (text, bold, color) in enumerate([
            (req, True, BLACK),
            (resp, True, resp_color),
            (detail, False, DARK_GRAY),
        ]):
            cell = out_table.cell(ri, ci)
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=30, bottom=30, left=80, right=80)
            p = cell.paragraphs[0]
            p.space_after = Pt(0)
            add_formatted_run(p, text, bold=bold, size=8, color=color)

    # ════════════════════════════════════════════════════════════
    # SECTION 5: RESPA
    # ════════════════════════════════════════════════════════════
    add_section_heading(doc, "RESPA Compliance (12 USC 2607)")

    add_body(doc,
        "Lead list purchases at a flat per-record price are explicitly permitted under RESPA Section 8(c)(2) "
        'as "payment for goods actually furnished." This is NOT a Marketing Services Agreement (CFPB Bulletin '
        "2015-05) and NOT a referral fee. There is no revenue sharing, no per-closed-loan pricing, and no "
        "reciprocal referral obligations.",
        size=9)

    # ════════════════════════════════════════════════════════════
    # SECTION 6: KEY STATUTES TABLE
    # ════════════════════════════════════════════════════════════
    add_section_heading(doc, "Governing Statutes & References")

    statutes = [
        ("FCRA", "15 USC 1681", "Consumer report definitions, CRA requirements"),
        ("TCPA", "47 USC 227", "Autodialer restrictions, DNC, cell phone rules"),
        ("CAN-SPAM", "15 USC 7701-7713", "Commercial email requirements"),
        ("RESPA Sec. 8", "12 USC 2607", "Prohibited kickbacks & referral fees"),
        ("RESPA 8(c)(2)", "12 USC 2607(c)(2)", "Safe harbor: payment for goods actually furnished"),
        ("FTC TSR", "16 CFR 310", "Telemarketing Sales Rule, DNC registry requirements"),
        ("HPPA (2026)", "Pub. L. 118-XXX", "Trigger lead ban -- credit bureau data only, not public records"),
        ("Feist v. Rural", "499 U.S. 340 (1991)", "Government facts are not copyrightable -- public records are free to aggregate"),
    ]

    stat_table = doc.add_table(rows=len(statutes) + 1, cols=3)
    stat_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ci, hdr in enumerate(["Law / Case", "Citation", "Relevance"]):
        cell = stat_table.cell(0, ci)
        set_cell_shading(cell, "0A2342")
        set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
        p = cell.paragraphs[0]
        p.space_after = Pt(0)
        add_formatted_run(p, hdr, bold=True, size=7.5, color=WHITE)

    for ri, (law, cite, rel) in enumerate(statutes, 1):
        bg = "F8F9FB" if ri % 2 == 0 else "FFFFFF"
        for ci, (text, bold, color) in enumerate([
            (law, True, BLACK),
            (cite, False, DARK_GRAY),
            (rel, False, DARK_GRAY),
        ]):
            cell = stat_table.cell(ri, ci)
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=20, bottom=20, left=80, right=80)
            p = cell.paragraphs[0]
            p.space_after = Pt(0)
            add_formatted_run(p, text, bold=bold, size=7.5, color=color)

    # ════════════════════════════════════════════════════════════
    # FOOTER
    # ════════════════════════════════════════════════════════════
    doc.add_paragraph()  # spacer

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.space_before = Pt(4)
    add_formatted_run(footer_p,
        "This document is a compliance summary, not legal advice. Consult a licensed attorney for formal guidance.",
        italic=True, size=7, color=MED_GRAY)

    footer_p2 = doc.add_paragraph()
    footer_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p2.space_before = Pt(2)
    add_formatted_run(footer_p2, "Confidential  |  Still Mind Creative  |  Updated March 2026",
                      size=7, color=MED_GRAY)

    # ── Save ──────────────────────────────────────────────────
    doc.save(str(output_path))
    print(f"Compliance one-pager: {output_path} ({output_path.stat().st_size / 1024:.0f} KB)")


def main():
    import argparse
    parser = argparse.ArgumentParser(description="Build compliance one-pager (Word doc)")
    parser.add_argument("--output", type=str, default=None)
    args = parser.parse_args()

    output_path = Path(args.output) if args.output else PROJECT_DIR / "data" / "workbooks" / "compliance_one_pager.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    build_compliance_doc(output_path)


if __name__ == "__main__":
    main()
